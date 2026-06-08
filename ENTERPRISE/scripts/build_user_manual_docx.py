#!/usr/bin/env python3
"""Generate docs/User_Manual.docx — run from repo root: python3 scripts/build_user_manual_docx.py"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "User_Manual.docx"
sys.path.insert(0, str(ROOT / "scripts"))
from user_manual_figures_data import MANUAL_FIGURES  # noqa: E402

FIG_DIR = ROOT / "docs" / "manual_figures"


def _shade_cell(cell, fill_hex: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def _add_figure_box(doc: Document, body: str) -> None:
    t = doc.add_table(rows=1, cols=1)
    t.autofit = True
    cell = t.rows[0].cells[0]
    _shade_cell(cell, "F1F5F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(body)
    run.italic = True
    run.font.size = Pt(10)
    for para in cell.paragraphs:
        para.paragraph_format.left_indent = Inches(0.08)


def _add_figure_image(doc: Document, figure_num: int) -> None:
    """Embed PNG from docs/manual_figures/ if present."""
    fname = None
    for num, fn, _, _ in MANUAL_FIGURES:
        if num == figure_num:
            fname = fn
            break
    path = FIG_DIR / fname if fname else None
    if path and path.is_file():
        doc.add_picture(str(path), width=Inches(6.25))
    else:
        _add_figure_box(
            doc,
            f"Figure {figure_num} — image missing. Run: python3 scripts/generate_manual_figure_images.py",
        )


def _h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def _bullet(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def _num(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Number")


def build() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Firmgate — User Manual")
    r.bold = True
    r.font.size = Pt(18)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ru = sub.add_run(
        "Your organisation name · registration details (optional)\n"
        "Company address · contact email · website"
    )
    ru.font.size = Pt(9)

    note = doc.add_paragraph()
    nn = note.add_run(
        "Figures below are branded placeholders generated from this repo. Right-click each image → Replace picture "
        "to drop in a real screenshot from your intranet (Full HD or similar). Redact passwords and secrets."
    )
    nn.italic = True
    nn.font.size = Pt(10)

    doc.add_page_break()

    _h(doc, "Contents", 1)
    toc = [
        "1. Introduction",
        "2. Signing in & browser tips",
        "3. Main navigation",
        "4. Home & news",
        "5. Blogs",
        "6. Events (calendar)",
        "7. Team Chat",
        "8. Workforce & dashboard",
        "9. Security clearance",
        "10. Security training",
        "11. Documents",
        "12. CRM",
        "13. About company",
        "14. Administration (authorised users)",
        "15. Troubleshooting",
        "16. Screenshot checklist",
    ]
    _num(doc, toc)

    doc.add_page_break()

    _h(doc, "1. Introduction", 1)
    _p(
        doc,
        "Firmgate is your internal portal for news, blogs, calendars, documents, "
        "collaboration, and company information. Access is controlled by your account and assigned permissions.",
    )
    _p(
        doc,
        "This manual describes common tasks. Exact labels may vary slightly if your administrator customises the portal.",
    )
    _add_figure_image(doc, 1)
    cap = doc.add_paragraph("Figure 1: Overview of the intranet after login.")
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)

    doc.add_page_break()
    _h(doc, "2. Signing in & browser tips", 1)
    _num(
        doc,
        [
            "Open your intranet URL (provided by IT — use the https:// address).",
            "Enter your email and password.",
            "Optional: use Stay signed in only on trusted work or personal devices.",
        ],
    )
    _p(
        doc,
        "Use a current browser (Chrome, Edge, Firefox, Safari). Office files may open in the browser or "
        "in OnlyOffice depending on configuration.",
    )
    _add_figure_image(doc, 2)
    c2 = doc.add_paragraph("Figure 2: Sign-in screen.")
    c2.runs[0].italic = True
    c2.runs[0].font.size = Pt(9)

    doc.add_page_break()
    _h(doc, "3. Main navigation", 1)
    _p(doc, "The top bar usually includes:")
    _bullet(
        doc,
        [
            "Logo / title — link to home.",
            "Search — search across allowed content (news, people, documents as configured).",
            "Main tabs — Home, Blogs, Events, Team Chat, Workforce, Workforce Dashboard, Security Clearance, "
            "Security Training, Documents, CRM, About Company, and possibly Administration if you have access.",
            "Your profile — account menu (top right).",
        ],
    )
    _add_figure_image(doc, 3)
    c3 = doc.add_paragraph("Figure 3: Main navigation.")
    c3.runs[0].italic = True
    c3.runs[0].font.size = Pt(9)

    _h(doc, "4. Home & news", 1)
    _p(
        doc,
        "Home typically shows announcements, news posts, or pinned updates. Click an article to read full content.",
    )
    _add_figure_image(doc, 4)
    doc.add_paragraph("Figure 4: Home / news.").runs[0].italic = True

    doc.add_page_break()
    _h(doc, "5. Blogs", 1)
    _p(doc, "Use Blogs for internal posts if enabled. You may read and, if permitted, create or comment.")
    _add_figure_image(doc, 5)
    doc.add_paragraph("Figure 5: Blogs.").runs[0].italic = True

    _h(doc, "6. Events (calendar)", 1)
    _p(doc, "Events opens the company calendar.")
    _bullet(
        doc,
        [
            "Switch between Day, Month, and Year views.",
            "Use arrows and Today to move around dates.",
            "If you may manage events, use Add event to create entries.",
            "Public holidays for configured regions may appear as read-only public holiday entries.",
        ],
    )
    _add_figure_image(doc, 6)
    doc.add_paragraph("Figure 6: Events calendar.").runs[0].italic = True

    doc.add_page_break()
    _h(doc, "7. Team Chat", 1)
    _p(doc, "Team Chat provides chat rooms and messages for staff collaboration (subject to your permissions).")
    _add_figure_image(doc, 7)
    doc.add_paragraph("Figure 7: Team Chat.").runs[0].italic = True

    _h(doc, "8. Workforce & dashboard", 1)
    _p(
        doc,
        "Workforce and Workforce Dashboard show tools or metrics your organisation enables (varies by deployment).",
    )
    _add_figure_image(doc, 8)
    doc.add_paragraph("Figure 8: Workforce / dashboard.").runs[0].italic = True

    _h(doc, "9. Security clearance", 1)
    _p(doc, "If enabled, this area supports security clearance workflows. Follow prompts from security or HR.")
    _add_figure_image(doc, 9)
    doc.add_paragraph("Figure 9: Security clearance.").runs[0].italic = True

    doc.add_page_break()
    _h(doc, "10. Security training", 1)
    _p(doc, "Security Training lists approved training files (videos and slide decks).")
    _num(
        doc,
        [
            "Select a file in the Training files list.",
            "Videos usually play in the main viewer.",
            "Presentations may open in an embedded viewer (e.g. OnlyOffice). If the viewer fails, contact IT.",
        ],
    )
    _add_figure_image(doc, 10)
    doc.add_paragraph("Figure 10: Security training.").runs[0].italic = True

    _h(doc, "11. Documents", 1)
    _h(doc, "11.1 Overview", 2)
    _p(
        doc,
        "Documents is the file browser. Most users see their own Home folder under All files, plus items "
        "shared with you or shared by you.",
    )
    _h(doc, "11.2 Uploading & folders", 2)
    _bullet(
        doc,
        [
            "Use Upload or drag-and-drop into the current folder (if allowed).",
            "Create folders where policy allows.",
            "Open a file to preview or download; Office formats may open in OnlyOffice when configured.",
        ],
    )
    _h(doc, "11.3 Sharing", 2)
    _p(
        doc,
        "To share with a colleague, use sharing on a file or folder you own (if enabled). Recipients need intranet accounts.",
    )
    _h(doc, "11.4 Personal & favourites", 2)
    _p(doc, "If available, mark items as favourites or personal per your organisation’s policy.")
    _add_figure_image(doc, 11)
    doc.add_paragraph("Figure 11: Documents browser.").runs[0].italic = True
    _add_figure_image(doc, 12)
    doc.add_paragraph("Figure 12: Upload / actions (optional).").runs[0].italic = True

    doc.add_page_break()
    _h(doc, "12. CRM", 1)
    _p(doc, "The CRM module stores leads, companies, or activities if your team uses it.")
    _add_figure_image(doc, 13)
    doc.add_paragraph("Figure 13: CRM.").runs[0].italic = True

    _h(doc, "13. About company", 1)
    _p(doc, "Company profile or reference information prepared by communications or leadership.")
    _add_figure_image(doc, 14)
    doc.add_paragraph("Figure 14: About company.").runs[0].italic = True

    doc.add_page_break()
    _h(doc, "14. Administration (authorised users)", 1)
    _p(
        doc,
        "Users with Administration access can manage users, roles, integrations (e.g. OnlyOffice), branding, "
        "backups, and other settings. This manual does not replace IT runbooks.",
    )
    _add_figure_image(doc, 15)
    doc.add_paragraph("Figure 15: Administration (redact secrets).").runs[0].italic = True

    _h(doc, "15. Troubleshooting", 1)
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Table Grid"
    hdr = ("Issue", "What to try")
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    rows = [
        ("Cannot sign in", "Reset password via IT; confirm correct URL and caps lock."),
        (
            "Office file / slides won’t open",
            "Try another browser; ensure HTTPS; ask IT to verify OnlyOffice URL, app URL, and JWT secret.",
        ),
        (
            "“Document security token” error",
            "IT must align JWT secret between Document Server and intranet Integrations.",
        ),
        ("Missing menu item", "Your role may not include that module — ask your administrator."),
    ]
    for ri, (a, b) in enumerate(rows, start=1):
        tbl.rows[ri].cells[0].text = a
        tbl.rows[ri].cells[1].text = b

    doc.add_page_break()
    _h(doc, "16. Screenshot checklist", 1)
    _p(doc, "Replace each placeholder image above with a real capture from your environment:")
    checklist = [
        "Signed-in home — full width.",
        "Sign-in page (redact credentials).",
        "Top navigation + search.",
        "News / home content.",
        "Blogs.",
        "Events month view.",
        "Team Chat.",
        "Workforce or dashboard.",
        "Security clearance.",
        "Security Training.",
        "Documents — root / Home.",
        "Documents — upload or share dialog (optional).",
        "CRM.",
        "About company.",
        "Administration sidebar (optional, no secrets).",
    ]
    _num(doc, checklist)

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run("Firmgate User Manual")
    fr.font.size = Pt(9)
    fr.font.italic = True

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
