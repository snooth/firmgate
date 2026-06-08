"""AI CV Builder — map a source CV into the user's Word template using an LLM."""

from __future__ import annotations

import json
import logging
import re
import urllib.error
import urllib.request
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any

from flask import current_app
from werkzeug.utils import secure_filename

from app.enterprise.ai_llm_settings import (
    PRODUCT_CV_BUILDER,
    iter_chat_completion_deltas,
    llm_api_key as _llm_api_key,
    llm_base_url as _llm_base_url,
    llm_configured as _llm_configured,
    llm_max_tokens as _llm_max_tokens,
    llm_model as _llm_model,
    llm_settings_public as _llm_settings_public,
    llm_temperature as _llm_temperature,
    _llm_ssl_context,
)
from app.enterprise.resource_pool_cv_import import extract_text_from_cv

log = logging.getLogger(__name__)


def llm_configured() -> bool:
    return _llm_configured(PRODUCT_CV_BUILDER)


def llm_settings_public() -> dict[str, Any]:
    return _llm_settings_public(PRODUCT_CV_BUILDER)

_MAX_BYTES = 12 * 1024 * 1024
_TEMPLATE_SUFFIX = ".docx"
_SOURCE_SUFFIXES = frozenset({".pdf", ".docx"})


def _user_dir(user_id: int) -> Path:
    root = Path(current_app.instance_path) / "ai_cv_builder" / str(int(user_id))
    root.mkdir(parents=True, exist_ok=True)
    return root


def _meta_path(user_id: int) -> Path:
    return _user_dir(user_id) / "meta.json"


def _load_meta(user_id: int) -> dict[str, Any]:
    p = _meta_path(user_id)
    if not p.is_file():
        return {}
    try:
        data = json.loads(p.read_text(encoding="utf-8"))
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def _save_meta(user_id: int, meta: dict[str, Any]) -> None:
    _meta_path(user_id).write_text(json.dumps(meta, indent=2), encoding="utf-8")


def _template_path(user_id: int) -> Path:
    return _user_dir(user_id) / "template.docx"


def _source_path(user_id: int) -> Path | None:
    meta = _load_meta(user_id)
    name = str(meta.get("source_filename") or "").strip()
    if not name:
        return None
    p = _user_dir(user_id) / name
    return p if p.is_file() else None


def _output_path(user_id: int) -> Path:
    return _user_dir(user_id) / "output.docx"


def public_status(user_id: int) -> dict[str, Any]:
    meta = _load_meta(user_id)
    tpl = _template_path(user_id)
    src = _source_path(user_id)
    out = _output_path(user_id)
    return {
        "template": {
            "set": tpl.is_file(),
            "name": meta.get("template_name") or (tpl.name if tpl.is_file() else ""),
            "updated_at": meta.get("template_updated_at"),
        },
        "source": {
            "set": src is not None,
            "name": meta.get("source_name") or "",
            "updated_at": meta.get("source_updated_at"),
        },
        "output": {
            "ready": out.is_file(),
            "built_at": meta.get("last_built_at"),
            "preview": meta.get("last_preview") or "",
        },
        "llm": llm_settings_public(),
    }


def save_template(user_id: int, filename: str, data: bytes) -> dict[str, Any]:
    if len(data) > _MAX_BYTES:
        raise ValueError("Template file is too large (max 12 MB).")
    name = secure_filename(filename) or "template.docx"
    if not name.lower().endswith(_TEMPLATE_SUFFIX):
        raise ValueError("CV template must be a Word document (.docx).")
    path = _template_path(user_id)
    path.write_bytes(data)
    meta = _load_meta(user_id)
    meta["template_name"] = name
    meta["template_updated_at"] = datetime.now(timezone.utc).isoformat()
    _save_meta(user_id, meta)
    return public_status(user_id)["template"]


def save_source(user_id: int, filename: str, data: bytes) -> dict[str, Any]:
    if len(data) > _MAX_BYTES:
        raise ValueError("Source CV is too large (max 12 MB).")
    name = secure_filename(filename) or "source.pdf"
    low = name.lower()
    if not any(low.endswith(s) for s in _SOURCE_SUFFIXES):
        raise ValueError("Source CV must be PDF or Word (.docx).")
    ext = Path(name).suffix.lower()
    stored = f"source{ext}"
    path = _user_dir(user_id) / stored
    path.write_bytes(data)
    meta = _load_meta(user_id)
    meta["source_name"] = name
    meta["source_filename"] = stored
    meta["source_updated_at"] = datetime.now(timezone.utc).isoformat()
    _save_meta(user_id, meta)
    return public_status(user_id)["source"]


def _extract_template_text(data: bytes) -> tuple[str, list[str]]:
    from docx import Document

    doc = Document(BytesIO(data))
    slots: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            slots.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = (p.text or "").strip()
                    if t:
                        slots.append(t)
    return "\n".join(slots), slots


def _collect_paragraph_objects(doc):
    """Yield paragraph objects in the same order as _extract_template_text."""
    for p in doc.paragraphs:
        if (p.text or "").strip():
            yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if (p.text or "").strip():
                        yield p


def _chat_completion_json(messages: list[dict[str, str]], *, max_tokens: int | None = None) -> str:
    from app.enterprise.ai_llm_http import effective_llm_api_key, openai_request_headers

    api_key = _llm_api_key(PRODUCT_CV_BUILDER)
    base = _llm_base_url(PRODUCT_CV_BUILDER)
    if not effective_llm_api_key(api_key, base):
        raise RuntimeError(
            "AI CV Builder is not configured. Set API key and base URL under Administration → AI Settings → AI CV Builder."
        )
    url = f"{base}/chat/completions"

    def _request(body: dict) -> str:
        req = urllib.request.Request(
            url,
            data=json.dumps(body).encode("utf-8"),
            headers=openai_request_headers(api_key, base),
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=180, context=_llm_ssl_context(PRODUCT_CV_BUILDER)) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        try:
            return (data["choices"][0]["message"]["content"] or "").strip()
        except (KeyError, IndexError, TypeError) as e:
            raise RuntimeError("Unexpected LLM response format") from e

    tokens = max_tokens if max_tokens is not None else _llm_max_tokens(PRODUCT_CV_BUILDER)
    base_body = {
        "model": _llm_model(PRODUCT_CV_BUILDER),
        "messages": messages,
        "temperature": _llm_temperature(PRODUCT_CV_BUILDER),
        "max_tokens": tokens,
    }
    try:
        return _request({**base_body, "response_format": {"type": "json_object"}})
    except urllib.error.HTTPError as e:
        if e.code in (400, 422):
            try:
                return _request(base_body)
            except urllib.error.HTTPError as e2:
                body_txt = e2.read().decode("utf-8", errors="replace")[:500]
                raise RuntimeError(f"LLM request failed ({e2.code}): {body_txt}") from e2
        body_txt = e.read().decode("utf-8", errors="replace")[:500]
        raise RuntimeError(f"LLM request failed ({e.code}): {body_txt}") from e
    except urllib.error.URLError as e:
        raise RuntimeError(f"Could not reach LLM API: {e.reason}") from e


def _parse_replacements(raw: str, expected: int) -> list[str]:
    try:
        obj = json.loads(raw)
    except json.JSONDecodeError:
        obj = {}
    reps = obj.get("replacements") if isinstance(obj, dict) else None
    if not isinstance(reps, list):
        reps = []
    out = [str(x or "").strip() for x in reps]
    if len(out) < expected:
        out.extend([""] * (expected - len(out)))
    return out[:expected]


def _apply_replacements_to_docx(template_bytes: bytes, replacements: list[str]) -> bytes:
    from docx import Document

    doc = Document(BytesIO(template_bytes))
    texts, _ = _extract_template_text(template_bytes)
    it = iter(replacements)
    for para in _collect_paragraph_objects(doc):
        try:
            new_text = next(it)
        except StopIteration:
            break
        if new_text:
            para.text = new_text
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_cv_messages(user_id: int, *, extra_instructions: str = "") -> tuple[list[dict[str, str]], bytes, list[str]]:
    tpl_path = _template_path(user_id)
    src_path = _source_path(user_id)
    if not tpl_path.is_file():
        raise ValueError("Upload your CV template (.docx) first.")
    if not src_path:
        raise ValueError("Upload the source CV (PDF or Word) first.")

    template_bytes = tpl_path.read_bytes()
    template_full, slots = _extract_template_text(template_bytes)
    source_bytes = src_path.read_bytes()
    source_text = extract_text_from_cv(str(_load_meta(user_id).get("source_name") or src_path.name), source_bytes)
    if not source_text.strip():
        raise ValueError("Could not extract text from the source CV.")
    if not slots:
        raise ValueError("The template has no text placeholders to fill. Add content to your .docx template.")

    slot_list = "\n".join(f"{i + 1}. {s[:500]}" for i, s in enumerate(slots))
    instr = (extra_instructions or "").strip()
    user_prompt = (
        f"You are an expert CV writer. Map the SOURCE CV content into the user's TEMPLATE structure.\n\n"
        f"TEMPLATE PARAGRAPHS ({len(slots)} slots — output exactly this many replacement strings in order):\n"
        f"{slot_list}\n\n"
        f"FULL TEMPLATE TEXT:\n{template_full[:12000]}\n\n"
        f"SOURCE CV:\n{source_text[:24000]}\n\n"
    )
    if instr:
        user_prompt += f"Additional instructions:\n{instr}\n\n"
    user_prompt += (
        "Return JSON only: {\"replacements\": [\"...\", ...], \"preview_markdown\": \"...\"} "
        f"with replacements array length exactly {len(slots)}. "
        "Each replacement is the new text for that template paragraph (plain text, no markdown). "
        "Preserve section headings from the template where they are labels; fill content paragraphs with "
        "accurate information from the source CV. Do not invent employers, dates, or credentials. "
        "preview_markdown is a readable markdown summary of the finished CV for on-screen preview."
    )
    messages = [
        {
            "role": "system",
            "content": "You transfer CV content into a fixed Word template. Output valid JSON only.",
        },
        {"role": "user", "content": user_prompt},
    ]
    return messages, template_bytes, slots


def iter_cv_build_deltas(messages: list[dict[str, str]]):
    return iter_chat_completion_deltas(PRODUCT_CV_BUILDER, messages, timeout=300)


def finalize_cv_build(user_id: int, raw: str, template_bytes: bytes, slots: list[str]) -> dict[str, Any]:
    replacements = _parse_replacements(raw, len(slots))
    try:
        preview = json.loads(raw).get("preview_markdown") or ""
    except Exception:
        preview = ""
    preview = str(preview).strip()

    out_bytes = _apply_replacements_to_docx(template_bytes, replacements)
    out_path = _output_path(user_id)
    out_path.write_bytes(out_bytes)

    meta = _load_meta(user_id)
    meta["last_built_at"] = datetime.now(timezone.utc).isoformat()
    meta["last_preview"] = preview[:20000] if preview else _preview_from_replacements(slots, replacements)
    _save_meta(user_id, meta)

    return {
        "preview_markdown": meta["last_preview"],
        "output_ready": True,
        "built_at": meta["last_built_at"],
        "slots_filled": len(slots),
    }


def build_cv(user_id: int, *, extra_instructions: str = "") -> dict[str, Any]:
    messages, template_bytes, slots = build_cv_messages(user_id, extra_instructions=extra_instructions)
    raw = _chat_completion_json(messages)
    return finalize_cv_build(user_id, raw, template_bytes, slots)


def _preview_from_replacements(slots: list[str], replacements: list[str]) -> str:
    parts = []
    for i, rep in enumerate(replacements):
        if not rep.strip():
            continue
        label = slots[i][:80] if i < len(slots) else f"Section {i + 1}"
        parts.append(f"### {label}\n\n{rep}\n")
    return "\n".join(parts).strip()


def output_file_path(user_id: int) -> Path | None:
    p = _output_path(user_id)
    return p if p.is_file() else None
