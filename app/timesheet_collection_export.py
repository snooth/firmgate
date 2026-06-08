"""Bulk export helpers for Timesheet Collection."""

from __future__ import annotations

from datetime import date, datetime, timezone
from io import BytesIO
import zipfile

from werkzeug.utils import secure_filename

from app.timesheet_signed import normalize_month, signed_file_path

_PDF_RENDER_DPI = 300
_TIMESHEET_MARGIN_IN = 0.2


def _unique_zip_name(base: str, used: set[str]) -> str:
    name = base or "timesheet.pdf"
    if name not in used:
        used.add(name)
        return name
    stem = name[:-4] if name.lower().endswith(".pdf") else name
    n = 2
    while True:
        candidate = secure_filename(f"{stem}_{n}.pdf") or f"timesheet_{n}.pdf"
        if candidate not in used:
            used.add(candidate)
            return candidate
        n += 1


def monthly_timesheets_docx_filename(month: str) -> str:
    month = normalize_month(month)
    y, m = month.split("-")
    month_name = date(int(y), int(m), 1).strftime("%B")
    return f"SolStak_Monthly_Timesheets_{month_name}_{y}.docx"


def monthly_timesheets_pdf_filename(month: str) -> str:
    month = normalize_month(month)
    y, m = month.split("-")
    month_name = date(int(y), int(m), 1).strftime("%B")
    return f"SolStak_Monthly_Timesheets_{month_name}_{y}.pdf"


def build_signed_timesheets_zip(month: str, rows: list[dict]) -> BytesIO | None:
    """ZIP of signed PDFs for uploaded rows in the given month."""
    month = normalize_month(month)
    uploaded = [r for r in rows if r.get("uploaded")]
    if not uploaded:
        return None

    buf = BytesIO()
    used_names: set[str] = set()
    added = 0
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for row in uploaded:
            user_id = row.get("user_id")
            if not user_id:
                continue
            path = signed_file_path(int(user_id), month)
            if not path or not path.exists():
                continue
            employee = secure_filename(str(row.get("employee_name") or f"user_{user_id}")) or f"user_{user_id}"
            original = secure_filename(str(row.get("original_name") or path.name)) or path.name
            if not original.lower().endswith(".pdf"):
                original = f"{original}.pdf"
            arcname = _unique_zip_name(f"{employee}_{original}", used_names)
            zf.write(path, arcname=arcname)
            added += 1
    if not added:
        return None
    buf.seek(0)
    return buf


def build_monthly_timesheets_docx(month: str, rows: list[dict]) -> BytesIO | None:
    """Single Word document containing rendered pages from all signed PDFs for the month."""
    try:
        import fitz  # pymupdf
    except ImportError as exc:
        raise RuntimeError("PDF export requires pymupdf. Install dependencies from requirements.txt.") from exc

    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.shared import Inches, Pt

    month = normalize_month(month)
    uploaded = sorted(
        [r for r in rows if r.get("uploaded")],
        key=lambda r: str(r.get("employee_name") or "").lower(),
    )
    if not uploaded:
        return None

    y, m = month.split("-")
    month_label = date(int(y), int(m), 1).strftime("%B %Y")

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    margin = Inches(_TIMESHEET_MARGIN_IN)
    section.left_margin = margin
    section.right_margin = margin
    section.top_margin = margin
    section.bottom_margin = margin
    target_width_in = (
        section.page_width.inches - section.left_margin.inches - section.right_margin.inches
    )
    target_height_in = (
        section.page_height.inches - section.top_margin.inches - section.bottom_margin.inches
    )
    render_zoom = _PDF_RENDER_DPI / 72.0

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    doc.add_heading(f"SolStak Monthly Timesheets — {month_label}", level=0)
    generated = datetime.now(timezone.utc).strftime("%d/%m/%Y %H:%M UTC")
    doc.add_paragraph(f"Compiled signed timesheets for {month_label}. Generated {generated}.")
    doc.add_page_break()

    added = 0
    for row in uploaded:
        user_id = row.get("user_id")
        if not user_id:
            continue
        path = signed_file_path(int(user_id), month)
        if not path or not path.exists():
            continue

        if added:
            doc.add_page_break()
        added += 1

        pdf = fitz.open(str(path))
        try:
            for page_index in range(len(pdf)):
                if page_index:
                    doc.add_page_break()
                page = pdf[page_index]
                page_width_in = page.rect.width / 72.0
                page_height_in = page.rect.height / 72.0
                fit_scale = min(
                    target_width_in / page_width_in if page_width_in else 1.0,
                    target_height_in / page_height_in if page_height_in else 1.0,
                )
                display_w = Inches(page_width_in * fit_scale)
                display_h = Inches(page_height_in * fit_scale)
                pix = page.get_pixmap(matrix=fitz.Matrix(render_zoom, render_zoom), alpha=False)
                img_buf = BytesIO(pix.tobytes("png"))
                doc.add_picture(img_buf, width=display_w, height=display_h)
        finally:
            pdf.close()

    if not added:
        return None

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_monthly_timesheets_pdf(month: str, rows: list[dict]) -> BytesIO | None:
    """Merge all signed timesheet PDFs for the month into one document (employee name order)."""
    try:
        import fitz  # pymupdf
    except ImportError as exc:
        raise RuntimeError("PDF export requires pymupdf. Install dependencies from requirements.txt.") from exc

    month = normalize_month(month)
    uploaded = sorted(
        [r for r in rows if r.get("uploaded")],
        key=lambda r: str(r.get("employee_name") or "").lower(),
    )
    if not uploaded:
        return None

    merged = fitz.open()
    added = 0
    try:
        for row in uploaded:
            user_id = row.get("user_id")
            if not user_id:
                continue
            path = signed_file_path(int(user_id), month)
            if not path or not path.exists():
                continue
            with fitz.open(str(path)) as src:
                merged.insert_pdf(src)
            added += 1
        if not added:
            return None
        buf = BytesIO()
        merged.save(buf)
        buf.seek(0)
        return buf
    finally:
        merged.close()
