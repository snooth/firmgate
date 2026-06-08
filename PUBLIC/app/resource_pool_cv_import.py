"""Extract resource fields from uploaded CV files (PDF / Word)."""

from __future__ import annotations

import re
from io import BytesIO
from typing import Any

_MAX_BYTES = 12 * 1024 * 1024
_ALLOWED_SUFFIXES = frozenset({".pdf", ".docx", ".doc"})

_EMAIL_RE = re.compile(
    r"\b[A-Za-z0-9][A-Za-z0-9._%+-]*@[A-Za-z0-9][A-Za-z0-9.-]*\.[A-Za-z]{2,}\b"
)
_PHONE_LIKE = re.compile(r"^[\d\s+().\-]{8,}$")
_URL_LIKE = re.compile(r"https?://|linkedin\.com|www\.", re.I)

_CLEARANCE_RULES: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"\bnegative\s+vett(?:ing|ed)\s*(?:level\s*)?3\b", re.I), "NV3"),
    (re.compile(r"\bnegative\s+vett(?:ing|ed)\s*(?:level\s*)?2\b", re.I), "NV2"),
    (re.compile(r"\bnegative\s+vett(?:ing|ed)\s*(?:level\s*)?1\b", re.I), "NV1"),
    (re.compile(r"\bNV[-\s]?3\b", re.I), "NV3"),
    (re.compile(r"\bNV[-\s]?2\b", re.I), "NV2"),
    (re.compile(r"\bNV[-\s]?1\b", re.I), "NV1"),
    (re.compile(r"\bTOP\s*SECRET\b", re.I), "TOP SECRET"),
    (re.compile(r"\bSECRET\b", re.I), "SECRET"),
    (re.compile(r"\bPROTECTED\b", re.I), "PROTECTED"),
    (re.compile(r"\bBASELINE\b", re.I), "Baseline"),
    (re.compile(r"\bBPSS\b", re.I), "BPSS"),
]

_LOCATION_RULES: list[tuple[str, re.Pattern[str]]] = [
    ("Canberra ACT", re.compile(r"\bCanberra(?:\s*,?\s*ACT)?\b", re.I)),
    ("Sydney NSW", re.compile(r"\bSydney(?:\s*,?\s*NSW)?\b", re.I)),
    ("Melbourne VIC", re.compile(r"\bMelbourne(?:\s*,?\s*VIC)?\b", re.I)),
    ("Brisbane QLD", re.compile(r"\bBrisbane(?:\s*,?\s*QLD)?\b", re.I)),
    ("Perth WA", re.compile(r"\bPerth(?:\s*,?\s*WA)?\b", re.I)),
    ("Adelaide SA", re.compile(r"\bAdelaide(?:\s*,?\s*SA)?\b", re.I)),
    ("Hobart TAS", re.compile(r"\bHobart(?:\s*,?\s*TAS)?\b", re.I)),
    ("Darwin NT", re.compile(r"\bDarwin(?:\s*,?\s*NT)?\b", re.I)),
]

_SKILL_SECTION_RE = re.compile(
    r"^(?:(?:key\s+)?skills?|technical\s+skills?|core\s+competenc(?:y|ies)|"
    r"technologies|tools?\s*&\s*technologies|competenc(?:y|ies))\s*:?\s*$",
    re.I,
)

_COMMON_SKILLS = (
    "AWS",
    "Azure",
    "GCP",
    "Google Cloud",
    "Kubernetes",
    "K8s",
    "Docker",
    "OpenShift",
    "Terraform",
    "Ansible",
    "Python",
    "Java",
    "JavaScript",
    "TypeScript",
    "React",
    "Angular",
    "Vue",
    "Node.js",
    "SQL",
    "PostgreSQL",
    "MySQL",
    "MongoDB",
    "Redis",
    "Kafka",
    "RabbitMQ",
    "Linux",
    "Windows Server",
    "VMware",
    "Citrix",
    "Active Directory",
    "PowerShell",
    "Bash",
    "Git",
    "CI/CD",
    "Jenkins",
    "GitLab",
    "GitHub Actions",
    "Agile",
    "Scrum",
    "ITIL",
    "ServiceNow",
    "Splunk",
    "ELK",
    "Prometheus",
    "Grafana",
    "Networking",
    "Firewall",
    "Cisco",
    "F5",
    "SAP",
    "Salesforce",
    "SharePoint",
    "Power BI",
    "Tableau",
    "Snowflake",
    "Databricks",
    "Spark",
    "Hadoop",
    "Machine Learning",
    "AI",
    "DevOps",
    "SRE",
    "Microservices",
    "REST",
    "API",
    "GraphQL",
    ".NET",
    "C#",
    "Go",
    "Golang",
    "Ruby",
    "PHP",
    "Spring",
    "Hibernate",
    "JUnit",
    "Selenium",
    "Jira",
    "Confluence",
)


def _s(v: Any) -> str:
    return ("" if v is None else str(v)).strip()


def allowed_cv_suffix(filename: str) -> str | None:
    name = _s(filename).lower()
    for suf in _ALLOWED_SUFFIXES:
        if name.endswith(suf):
            return suf
    return None


def extract_text_from_cv(filename: str, data: bytes) -> str:
    """Return plain text from PDF or Word CV bytes."""
    if len(data) > _MAX_BYTES:
        raise ValueError("file too large (max 12 MB)")
    suf = allowed_cv_suffix(filename)
    if not suf:
        raise ValueError("unsupported file type (PDF or Word required)")
    if suf == ".pdf":
        return _pdf_text(data)
    if suf == ".docx":
        return _docx_text(data)
    raise ValueError(".doc files are not supported — save as .docx or PDF")


def _pdf_text(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise RuntimeError("PDF support is not installed") from e
    reader = PdfReader(BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        if t.strip():
            parts.append(t)
    return "\n".join(parts)


def _docx_text(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("Word support is not installed") from e
    doc = Document(BytesIO(data))
    parts = [p.text for p in doc.paragraphs if _s(p.text)]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = _s(cell.text)
                if t:
                    parts.append(t)
    return "\n".join(parts)


def _split_name(full: str) -> tuple[str, str]:
    parts = [p for p in full.split() if p]
    if len(parts) >= 2:
        return (" ".join(parts[:-1])[:120], parts[-1][:120])
    if parts:
        return (parts[0][:120], "")
    return ("", "")


def _line_looks_like_name(line: str) -> bool:
    s = _s(line)
    if not s or len(s) > 80:
        return False
    if _EMAIL_RE.search(s) or _URL_LIKE.search(s) or _PHONE_LIKE.match(s):
        return False
    if re.search(r"\d{3,}", s):
        return False
    low = s.lower()
    if any(
        w in low
        for w in (
            "curriculum",
            "resume",
            "résumé",
            "profile",
            "summary",
            "experience",
            "education",
            "skills",
            "clearance",
            "phone",
            "mobile",
            "email",
            "address",
        )
    ):
        return False
    words = s.split()
    if len(words) < 2 or len(words) > 5:
        return False
    alpha = sum(1 for w in words if w.replace("-", "").replace("'", "").isalpha())
    return alpha >= len(words) - 1


def _extract_name(lines: list[str], text: str) -> tuple[str, str, str]:
    for pat in (
        re.compile(r"(?:full\s+)?name\s*[:\-]\s*(.+)", re.I),
        re.compile(r"^name\s*[:\-]\s*(.+)", re.I | re.M),
    ):
        m = pat.search(text)
        if m:
            full = _s(m.group(1).split("\n")[0])
            if _line_looks_like_name(full):
                g, f = _split_name(full)
                return (g, f, full[:255])

    for line in lines[:12]:
        if _line_looks_like_name(line):
            full = _s(line)
            g, f = _split_name(full)
            return (g, f, full[:255])
    return ("", "", "")


def _extract_email(text: str) -> str:
    emails = _EMAIL_RE.findall(text)
    for raw in emails:
        e = raw.lower()
        if any(x in e for x in ("example.com", "email.com", "domain.com")):
            continue
        return raw[:255]
    return ""


def _extract_clearance(text: str) -> str:
    for rx, label in _CLEARANCE_RULES:
        if rx.search(text):
            return label
    return ""


def _extract_location(text: str) -> str:
    for label, rx in _LOCATION_RULES:
        if rx.search(text):
            return label
    m = re.search(
        r"(?:location|based\s+in|address)\s*[:\-]\s*([^\n]{3,80})",
        text,
        re.I,
    )
    if m:
        loc = _s(m.group(1))
        if loc and not _EMAIL_RE.search(loc):
            return loc[:255]
    return ""


def _skills_from_section(lines: list[str]) -> list[str]:
    out: list[str] = []
    in_section = False
    for line in lines:
        s = _s(line)
        if not s:
            if in_section and out:
                break
            continue
        if _SKILL_SECTION_RE.match(s):
            in_section = True
            continue
        if in_section:
            if re.match(
                r"^(?:experience|employment|work\s+history|education|certification|"
                r"projects|clearance|profile|summary)\b",
                s,
                re.I,
            ):
                break
            for part in re.split(r"[,;|•·\t/]", s):
                t = _s(part).strip("•·- ")
                if 2 <= len(t) <= 80 and not t.endswith(":"):
                    out.append(t)
    return out


def _skills_from_keywords(text: str, known: list[str]) -> list[str]:
    found: list[str] = []
    hay = text.lower()
    candidates = list(dict.fromkeys([*known, *_COMMON_SKILLS]))
    for sk in candidates:
        t = _s(sk)
        if not t or len(t) < 2:
            continue
        if re.search(rf"\b{re.escape(t.lower())}\b", hay):
            found.append(t)
    return found


def parse_cv_text(text: str, *, known_skills: list[str] | None = None) -> dict[str, Any]:
    """Heuristic parse of CV plain text into resource pool fields."""
    raw = (text or "").replace("\r", "\n")
    lines = [_s(ln) for ln in raw.split("\n") if _s(ln)]
    joined = "\n".join(lines)

    given, family, full = _extract_name(lines, joined)
    email = _extract_email(joined)
    clearance = _extract_clearance(joined)
    location = _extract_location(joined)

    section_skills = _skills_from_section(lines)
    keyword_skills = _skills_from_keywords(joined, known_skills or [])
    skills: list[str] = []
    seen: set[str] = set()
    for sk in section_skills + keyword_skills:
        key = sk.lower()
        if key not in seen:
            seen.add(key)
            skills.append(sk[:80])
        if len(skills) >= 24:
            break

    warnings: list[str] = []
    if not full and not (given or family):
        warnings.append("Could not detect a name — enter it before saving.")
    if not email:
        warnings.append("No email found in the CV.")
    if not skills:
        warnings.append("No skills detected — add them manually if needed.")
    if not clearance:
        warnings.append("No clearance level detected.")
    if not location:
        warnings.append("No location detected.")

    return {
        "given_name": given,
        "family_name": family,
        "full_name": full or f"{given} {family}".strip(),
        "email": email,
        "skills": skills,
        "clearance_level": clearance,
        "location": location,
        "warnings": warnings,
    }


def parse_cv_file(filename: str, data: bytes, *, known_skills: list[str] | None = None) -> dict[str, Any]:
    text = extract_text_from_cv(filename, data)
    if not _s(text):
        raise ValueError("could not read text from this file — try a text-based PDF or .docx")
    parsed = parse_cv_text(text, known_skills=known_skills)
    parsed["text_length"] = len(text)
    return parsed
